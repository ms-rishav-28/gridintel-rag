"""
document_processor.py - Parse, chunk, embed, and ingest documents into LanceDB.

Parsing tiers:
  PDF text    -> pymupdf4llm (primary) -> pdfplumber table fallback -> pytesseract OCR
  DOCX        -> python-docx (paragraphs + tables as markdown)
  Webpage     -> web_ingestion.py (separate module)
"""

# CODEX-FIX: replace simple LangChain loaders with tiered parsing, vision chunks, and LanceDB ingestion.

import asyncio
import hashlib
import io
import logging
import uuid
from enum import Enum
from pathlib import Path

from langchain_text_splitters import MarkdownHeaderTextSplitter, RecursiveCharacterTextSplitter

from app.core.config import get_settings
from app.services.convex_service import get_convex_service
from app.services.embedding_service import get_embedding_service
from app.services.vector_store import ChunkRecord, get_vector_store
from app.services.vision_service import get_vision_service

logger = logging.getLogger(__name__)


class PDFParseStrategy(str, Enum):
    FAST = "fast"
    THOROUGH = "thorough"
    OCR = "ocr"


class DocumentProcessor:
    HEADERS_TO_SPLIT = [("#", "h1"), ("##", "h2"), ("###", "h3")]
    CHUNK_SIZE = 1500
    CHUNK_OVERLAP = 200

    def __init__(self):
        self._settings = get_settings()
        self._splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.CHUNK_SIZE,
            chunk_overlap=self.CHUNK_OVERLAP,
            length_function=len,
        )
        self._header_splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=self.HEADERS_TO_SPLIT,
            strip_headers=False,
        )

    async def ingest_file(
        self,
        file_bytes: bytes,
        filename: str,
        doc_id: str,
        job_id: str,
        storage_id: str | None = None,
    ) -> dict[str, int | str | None]:
        """
        Main entry point for file ingestion.
        Returns { chunk_count, image_count, error }.
        Updates Convex job and document status throughout.
        """
        convex = get_convex_service()
        ext = Path(filename).suffix.lower()
        source_type = self._source_type_from_extension(ext)

        try:
            await convex.update_job(
                job_id,
                status="processing",
                progress_message="Parsing document...",
            )
            await convex.update_document(
                doc_id,
                ingestion_status="processing",
                storage_id=storage_id,
            )

            if ext == ".pdf":
                text_md, images = await self._parse_pdf(file_bytes)
            elif ext in (".docx", ".doc"):
                text_md, images = await self._parse_docx(file_bytes)
            elif ext == ".txt":
                text_md = file_bytes.decode("utf-8", errors="replace")
                images = []
            else:
                raise ValueError(f"Unsupported file type: {ext}")

            await convex.update_job(job_id, progress_message="Chunking and embedding...")
            text_chunks = await self._chunk_and_embed(
                text_md,
                doc_id,
                filename,
                source_type=source_type,
            )

            image_chunks: list[ChunkRecord] = []
            if images and self._settings.ENABLE_VISION:
                await convex.update_job(
                    job_id,
                    progress_message=f"Describing {len(images)} images...",
                    total_chunks=len(text_chunks) + len(images),
                )
                image_chunks = await self._process_images(images, doc_id, filename, source_type)

            all_chunks = text_chunks + image_chunks
            await convex.update_job(job_id, progress_message="Indexing into vector store...")
            await get_vector_store().add_chunks(all_chunks)

            await convex.update_document(
                doc_id,
                ingestion_status="done",
                chunk_count=len(all_chunks),
                image_count=len(image_chunks),
            )
            await convex.update_job(
                job_id,
                status="done",
                progress_message="Done",
                total_chunks=len(all_chunks),
                processed_chunks=len(all_chunks),
            )
            logger.info(
                "Ingested %s: %s text chunks, %s image chunks",
                filename,
                len(text_chunks),
                len(image_chunks),
            )
            return {
                "chunk_count": len(all_chunks),
                "image_count": len(image_chunks),
                "error": None,
            }

        except Exception as exc:
            logger.error("Document ingestion failed for %s: %s", filename, exc, exc_info=True)
            await convex.update_document(
                doc_id,
                ingestion_status="failed",
                error_message=str(exc),
            )
            await convex.update_job(job_id, status="failed", error_message=str(exc))
            return {"chunk_count": 0, "image_count": 0, "error": str(exc)}

    def _source_type_from_extension(self, ext: str) -> str:
        if ext == ".pdf":
            return "pdf"
        if ext in (".docx", ".doc"):
            return "docx"
        if ext == ".txt":
            return "txt"
        return ext.lstrip(".")

    # -- PDF parsing ------------------------------------------------------------

    async def _parse_pdf(self, pdf_bytes: bytes) -> tuple[str, list[dict]]:
        return await asyncio.to_thread(self._parse_pdf_sync, pdf_bytes)

    def _parse_pdf_sync(self, pdf_bytes: bytes) -> tuple[str, list[dict]]:
        import fitz

        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        try:
            total_chars = sum(len(page.get_text()) for page in doc)
            avg_chars_per_page = total_chars / max(len(doc), 1)
            strategy = (
                PDFParseStrategy.OCR
                if avg_chars_per_page < 100
                else PDFParseStrategy.THOROUGH
            )
            logger.info(
                "PDF parse strategy selected: %s (%s chars/page avg, %s pages)",
                strategy.value,
                round(avg_chars_per_page),
                len(doc),
            )

            if strategy == PDFParseStrategy.OCR:
                text_md = self._ocr_pdf_with_pymupdf(pdf_bytes)
            else:
                import pymupdf4llm

                text_md = pymupdf4llm.to_markdown(doc)
                try:
                    text_md = self._merge_pdfplumber_tables(pdf_bytes, text_md)
                except Exception as exc:
                    logger.warning("pdfplumber table merge failed: %s", exc)

            images: list[dict] = []
            for page_num, page in enumerate(doc, start=1):
                page_text = page.get_text()[:1000]
                for img_index, img in enumerate(page.get_images(full=True)):
                    xref = img[0]
                    try:
                        img_data = doc.extract_image(xref)
                        if img_data and img_data.get("image"):
                            images.append(
                                {
                                    "bytes": img_data["image"],
                                    "page_number": page_num,
                                    "image_index": img_index,
                                    "context_text": page_text,
                                }
                            )
                    except Exception:
                        continue

            return text_md, images
        finally:
            doc.close()

    def _ocr_pdf_with_pymupdf(self, pdf_bytes: bytes) -> str:
        """OCR pages one at a time to avoid loading large scanned PDFs into memory."""
        try:
            import fitz
            import pytesseract
            from PIL import Image

            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            pages: list[str] = []
            try:
                for page_index, page in enumerate(doc, start=1):
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                    image = Image.open(io.BytesIO(pix.tobytes("png")))
                    text = pytesseract.image_to_string(image, lang="eng")
                    pages.append(f"## Page {page_index}\n\n{text}")
            finally:
                doc.close()
            return "\n\n".join(pages)
        except Exception as exc:
            logger.warning("PDF OCR failed: %s", exc)
            return ""

    def _merge_pdfplumber_tables(self, pdf_bytes: bytes, base_md: str) -> str:
        """Extract tables via pdfplumber and append as GitHub-flavored markdown."""
        import pdfplumber

        table_sections: list[str] = []
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                tables = page.extract_tables()
                for table_index, table in enumerate(tables):
                    if not table or not table[0]:
                        continue
                    header = " | ".join(str(cell or "") for cell in table[0])
                    sep = " | ".join(["---"] * len(table[0]))
                    rows = [
                        " | ".join(str(cell or "") for cell in row)
                        for row in table[1:]
                    ]
                    table_md = f"\n### Table {table_index + 1} (page {page_num})\n\n"
                    table_md += f"| {header} |\n| {sep} |\n"
                    table_md += "\n".join(f"| {row} |" for row in rows)
                    table_sections.append(table_md)
        if table_sections:
            return base_md + "\n\n## Extracted Tables\n\n" + "\n\n".join(table_sections)
        return base_md

    # -- DOCX parsing -----------------------------------------------------------

    async def _parse_docx(self, docx_bytes: bytes) -> tuple[str, list[dict]]:
        return await asyncio.to_thread(self._parse_docx_sync, docx_bytes)

    def _parse_docx_sync(self, docx_bytes: bytes) -> tuple[str, list[dict]]:
        import docx

        doc = docx.Document(io.BytesIO(docx_bytes))
        parts: list[str] = []
        heading_map = {
            "Heading 1": "#",
            "Heading 2": "##",
            "Heading 3": "###",
            "Heading 4": "####",
        }

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            prefix = heading_map.get(style_name, "")
            text = para.text.strip()
            if not text:
                continue
            parts.append(f"{prefix} {text}" if prefix else text)

        for table in doc.tables:
            rows_md: list[str] = []
            for row_index, row in enumerate(table.rows):
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                rows_md.append("| " + " | ".join(cells) + " |")
                if row_index == 0:
                    rows_md.append("| " + " | ".join(["---"] * len(cells)) + " |")
            parts.append("\n".join(rows_md))

        images: list[dict] = []
        image_index = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    images.append(
                        {
                            "bytes": rel.target_part.blob,
                            "page_number": 1,
                            "image_index": image_index,
                            "context_text": "",
                        }
                    )
                    image_index += 1
                except Exception:
                    continue

        return "\n\n".join(parts), images

    # -- Chunking + embedding ---------------------------------------------------

    async def _chunk_and_embed(
        self,
        text_md: str,
        doc_id: str,
        doc_name: str,
        source_type: str,
        source_url: str | None = None,
    ) -> list[ChunkRecord]:
        if not text_md.strip():
            logger.warning("Empty text extracted for %s", doc_name)
            return []

        header_chunks = self._header_splitter.split_text(text_md)
        raw_chunks: list[dict[str, str | None]] = []
        for header_chunk in header_chunks:
            sub_chunks = self._splitter.split_text(header_chunk.page_content)
            for sub_chunk in sub_chunks:
                heading = " > ".join(
                    value for _, value in sorted(header_chunk.metadata.items()) if value
                ) or None
                raw_chunks.append({"text": sub_chunk, "heading": heading})

        if not raw_chunks:
            for sub_chunk in self._splitter.split_text(text_md):
                raw_chunks.append({"text": sub_chunk, "heading": None})

        texts = [str(chunk["text"]) for chunk in raw_chunks]
        vectors = await get_embedding_service().async_encode_dense(texts)

        records: list[ChunkRecord] = []
        for index, (chunk, vector) in enumerate(zip(raw_chunks, vectors)):
            records.append(
                ChunkRecord(
                    chunk_id=str(uuid.uuid4()),
                    doc_id=doc_id,
                    doc_name=doc_name,
                    source_type=source_type,
                    source_url=source_url,
                    page_number=None,
                    chunk_index=index,
                    section_heading=chunk["heading"],
                    chunk_type="text",
                    content=str(chunk["text"]),
                    vector=vector,
                )
            )
        return records

    # -- Image processing -------------------------------------------------------

    async def _process_images(
        self,
        images: list[dict],
        doc_id: str,
        doc_name: str,
        source_type: str,
    ) -> list[ChunkRecord]:
        vision = get_vision_service()
        embedder = get_embedding_service()
        records: list[ChunkRecord] = []

        for image in images:
            description = await vision.describe_image(
                image_bytes=image["bytes"],
                page_number=image["page_number"],
                image_index=image["image_index"],
                context_text=image.get("context_text", ""),
            )
            if not description:
                continue
            vector = (await embedder.async_encode_dense([description]))[0]
            records.append(
                ChunkRecord(
                    chunk_id=str(uuid.uuid4()),
                    doc_id=doc_id,
                    doc_name=doc_name,
                    source_type=source_type,
                    source_url=None,
                    page_number=image["page_number"],
                    chunk_index=image["image_index"],
                    section_heading=None,
                    chunk_type="image_description",
                    content=description,
                    vector=vector,
                )
            )

        return records

    @staticmethod
    def compute_sha256(file_bytes: bytes) -> str:
        return hashlib.sha256(file_bytes).hexdigest()


_processor: DocumentProcessor | None = None


def get_document_processor() -> DocumentProcessor:
    global _processor
    if _processor is None:
        _processor = DocumentProcessor()
    return _processor
